import os
import asyncio
import pypdf
import docx
from app.core.exceptions import FileValidationError
import logging

logger = logging.getLogger("finragvault.services.parser")


class ParserService:
    """Synchronous & Asynchronous service to parse text contents from supported document types."""

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """Parses a PDF document, extracting raw text page-by-page."""
        try:
            reader = pypdf.PdfReader(file_path)
            text_blocks = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_blocks.append(page_text)
            return "\n\n".join(text_blocks)
        except Exception as exc:
            logger.error(f"Error parsing PDF '{file_path}': {str(exc)}", exc_info=True)
            raise FileValidationError(f"Could not parse PDF document content: {str(exc)}")

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """Parses a DOCX document, extracting paragraph and table text blocks."""
        try:
            doc = docx.Document(file_path)
            text_blocks = []
            
            # Extract paragraph text
            for para in doc.paragraphs:
                if para.text.strip():
                    text_blocks.append(para.text)
                    
            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_blocks.append(" | ".join(row_text))
                        
            return "\n\n".join(text_blocks)
        except Exception as exc:
            logger.error(f"Error parsing DOCX '{file_path}': {str(exc)}", exc_info=True)
            raise FileValidationError(f"Could not parse DOCX document content: {str(exc)}")

    @staticmethod
    def _parse_txt(file_path: str) -> str:
        """Parses a TXT document using UTF-8 decoding."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        except Exception as exc:
            logger.error(f"Error parsing TXT '{file_path}': {str(exc)}", exc_info=True)
            raise FileValidationError(f"Could not parse TXT document content: {str(exc)}")

    @classmethod
    async def parse_document_async(cls, file_path: str, extension: str) -> str:
        """Runs the file parser in a background thread to prevent blocking the event loop.
        
        Args:
            file_path (str): Path to local file.
            extension (str): File extension without dot, e.g. "pdf".
            
        Returns:
            str: Full extracted text content.
        """
        ext = extension.lower().strip()
        
        if ext == "pdf":
            return await asyncio.to_thread(cls._parse_pdf, file_path)
        elif ext == "docx":
            return await asyncio.to_thread(cls._parse_docx, file_path)
        elif ext in ["txt", "text"]:
            return await asyncio.to_thread(cls._parse_txt, file_path)
        else:
            raise FileValidationError(f"Unsupported file format extension for parsing: '.{extension}'")
